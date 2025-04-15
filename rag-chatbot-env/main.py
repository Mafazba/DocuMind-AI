import os
import tempfile
import logging
from typing import Tuple, List, Dict, Optional
from datetime import datetime, timedelta
import time
from pathlib import Path
import uuid
import asyncio
import pickle
import shutil
from fastapi import FastAPI, HTTPException, UploadFile, File
import numpy as np
import faiss
import pdfplumber
from docx import Document
import pandas as pd
from langchain.text_splitter import CharacterTextSplitter
from sentence_transformers import SentenceTransformer
from presidio_analyzer import AnalyzerEngine
from presidio_anonymizer import AnonymizerEngine
from pydantic import BaseModel
from transformers import MarianMTModel, MarianTokenizer, BartForConditionalGeneration, BartTokenizer
from langdetect import detect
from rouge_score import rouge_scorer
import spacy
import ollama

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = FastAPI()

# Model management singleton
class ModelManager:
    _instance = None
    _models = {}

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._initialize()
        return cls._instance

    def _initialize(self):
        self._models = {
            "spacy": None,
            "sentence_transformer": None,
            "translation_en": None,
            "translation_ar": None,
            "summarizer": None
        }

    def get_spacy(self) -> Optional[spacy.language.Language]:
        if self._models["spacy"] is None:
            try:
                self._models["spacy"] = spacy.load("en_core_web_sm", disable=["parser"])
                logger.debug("Loaded spaCy model")
            except Exception as e:
                logger.error(f"Failed to load spaCy: {e}")
        return self._models["spacy"]

    def get_sentence_transformer(self) -> Optional[SentenceTransformer]:
        if self._models["sentence_transformer"] is None:
            try:
                self._models["sentence_transformer"] = SentenceTransformer(
                    'nomic-ai/nomic-embed-text-v1', trust_remote_code=True
                )
                logger.debug("Loaded SentenceTransformer")
            except Exception as e:
                logger.error(f"Failed to load SentenceTransformer: {e}")
        return self._models["sentence_transformer"]

    def get_translation(self, lang: str) -> Optional[tuple]:
        lang = lang.lower()
        if lang not in ["en", "ar"]:
            return None
        key = f"translation_{lang}"
        if self._models[key] is None:
            try:
                model_name = "Helsinki-NLP/opus-mt-mul-en" if lang == "en" else "Helsinki-NLP/opus-mt-en-ar"
                tokenizer = MarianTokenizer.from_pretrained(model_name)
                model = MarianMTModel.from_pretrained(model_name)
                self._models[key] = (tokenizer, model)
                logger.debug(f"Loaded translation model for {lang}")
            except Exception as e:
                logger.error(f"Failed to load translation model for {lang}: {e}")
        return self._models[key]

    def get_summarizer(self) -> Optional[tuple]:
        if self._models["summarizer"] is None:
            try:
                tokenizer = BartTokenizer.from_pretrained("facebook/bart-large-cnn")
                model = BartForConditionalGeneration.from_pretrained("facebook/bart-large-cnn")
                self._models["summarizer"] = (tokenizer, model)
                logger.debug("Loaded summarizer")
            except Exception as e:
                logger.error(f"Failed to load summarizer: {e}")
        return self._models["summarizer"]

# Initialize PII redaction
analyzer = AnalyzerEngine()
anonymizer = AnonymizerEngine()

# Session and FAISS cache
SESSION_HISTORY: Dict[str, List[Tuple[str, str, datetime]]] = {}
HISTORY_LIMIT = 3
SESSION_TIMEOUT = timedelta(minutes=30)
FAISS_CACHE = {"index": None, "metadata": None, "path": None}

class QueryRequest(BaseModel):
    query: str
    top_k: int = 5
    session_id: Optional[str] = None
    target_language: str = "en" #en or ar 
    summarize: Optional[str] = None   #context or answer or None

def cleanup_sessions():
    """Remove expired sessions."""
    current_time = datetime.utcnow()
    for sid in list(SESSION_HISTORY.keys()):
        if SESSION_HISTORY[sid] and current_time - SESSION_HISTORY[sid][-1][2] > SESSION_TIMEOUT:
            logger.debug(f"Cleaning session: {sid}")
            del SESSION_HISTORY[sid]

def redact_pii(text: str) -> str:
    """Redact PII using Presidio."""
    if not text or not isinstance(text, str):
        return text
    try:
        results = analyzer.analyze(text=text, entities=["PERSON", "PHONE_NUMBER", "EMAIL_ADDRESS", "CREDIT_CARD", "IBAN"], language="en")
        redacted = anonymizer.anonymize(text=text, analyzer_results=results)
        return redacted.text
    except Exception as e:
        logger.error(f"PII redaction error: {e}")
        return text

def read_file_content(file_path: str) -> Tuple[str, List[int], Dict[str, any]]:
    """Read and redact content from PDF, DOCX, or XLSX."""
    ext = Path(file_path).suffix.lower().lstrip(".")
    error_report = {"file": file_path, "errors": []}
    try:
        if ext == "pdf":
            with pdfplumber.open(file_path) as pdf:
                text = []
                page_numbers = []
                for i, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text(layout=True) or ""
                    if page_text:
                        text.append(redact_pii(page_text))
                    for table in page.extract_tables() or []:
                        try:
                            table_text = "\n".join([",".join(str(cell) if cell else "" for cell in row) for row in table])
                            if table_text.strip():
                                text.append(redact_pii(table_text))
                        except Exception as e:
                            error_report["errors"].append(f"Table error on page {i}: {str(e)}")
                    page_numbers.extend([i] * max(1, len(page_text.split("\n"))))
                full_text = "\n".join(text)
                if not full_text.strip():
                    error_report["errors"].append("No content extracted")
                return full_text, page_numbers, error_report
        elif ext == "docx":
            doc = Document(file_path)
            text = [redact_pii(p.text) for p in doc.paragraphs if p.text.strip()]
            for i, table in enumerate(doc.tables):
                try:
                    for row in table.rows:
                        row_text = "\t".join(cell.text for cell in row.cells if cell.text.strip())
                        if row_text:
                            text.append(redact_pii(row_text))
                except Exception as e:
                    error_report["errors"].append(f"Table {i+1} error: {str(e)}")
            full_text = "\n".join(text)
            if not full_text.strip():
                error_report["errors"].append("No content extracted")
            return full_text, [1] * max(1, len(full_text.split("\n"))), error_report
        elif ext == "xlsx":
            xls = pd.ExcelFile(file_path)
            text = []
            for sheet in xls.sheet_names:
                try:
                    df = pd.read_excel(xls, sheet_name=sheet, dtype=str).fillna("")
                    sheet_text = df.to_string(index=False, header=True)
                    if sheet_text.strip():
                        text.append(redact_pii(sheet_text))
                except Exception as e:
                    error_report["errors"].append(f"Sheet error: {str(e)}")
            full_text = "\n".join(text)
            if not full_text.strip():
                error_report["errors"].append("No content extracted")
            return full_text, [1] * max(1, len(full_text.split("\n"))), error_report
        error_report["errors"].append(f"Unsupported extension: {ext}")
        return "", [], error_report
    except Exception as e:
        error_report["errors"].append(f"Unexpected error: {str(e)}")
        return "", [], error_report

def chunk_text(text_data: Tuple[str, List[int]], chunk_size: int = 512, chunk_overlap: int = 50) -> Tuple[List[str], List[int]]:
    """Split text into chunks."""
    text, page_numbers = text_data
    if not text:
        return [], []
    try:
        splitter = CharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separator="\n",
            keep_separator=True
        )
        chunks = splitter.split_text(text.strip())
        if not chunks:
            return [], []
        chunk_page_numbers = []
        line_starts = [0]
        pos = 0
        for line in text.split("\n"):
            pos += len(line) + 1
            line_starts.append(pos)
        for chunk in chunks:
            start = text.find(chunk)
            idx = min(range(len(line_starts)), key=lambda i: abs(line_starts[i] - start))
            page = page_numbers[idx] if idx < len(page_numbers) else page_numbers[-1] if page_numbers else 1
            chunk_page_numbers.append(page)
        return chunks, chunk_page_numbers
    except Exception as e:
        logger.error(f"Chunking error: {e}")
        return [], []

def migrate_root_files(storage_dir: str) -> None:
    """Move FAISS files from root to storage_dir."""
    storage_path = Path(storage_dir).resolve()
    storage_path.mkdir(parents=True, exist_ok=True)
    for file_name in ["faiss_index.bin", "faiss_metadata.pkl"]:
        root_file = Path(file_name)
        target_file = storage_path / file_name
        if root_file.exists() and not target_file.exists():
            try:
                shutil.move(str(root_file), str(target_file))
                logger.info(f"Moved {file_name} to {target_file}")
            except Exception as e:
                logger.error(f"Failed to move {file_name}: {e}")

async def embed_with_timeout(chunks: List[str], model: SentenceTransformer, timeout: float = 180.0) -> np.ndarray:
    """Embed chunks with timeout."""
    try:
        batch_size = min(64, max(8, len(chunks) // 4))  # Dynamic batch size
        embeddings = await asyncio.wait_for(
            asyncio.get_event_loop().run_in_executor(
                None,
                lambda: model.encode(chunks, batch_size=batch_size, normalize_embeddings=True, show_progress_bar=False)
            ),
            timeout=timeout
        )
        return np.array(embeddings, dtype=np.float32)
    except asyncio.TimeoutError:
        logger.error("Embedding timeout")
        raise HTTPException(status_code=504, detail="Embedding timeout")
    except Exception as e:
        logger.error(f"Embedding error: {e}")
        raise

async def embed_and_store_chunks(
    chunks: List[str],
    file_name: str,
    chunk_page_numbers: List[int],
    index_file: str = "faiss_index.bin",
    meta_file: str = "faiss_metadata.pkl",
    storage_dir: str = os.getenv("STORAGE_DIR", "data")
) -> Tuple[List[str], float]:
    """Embed and store chunks in FAISS."""
    if not chunks or len(chunks) != len(chunk_page_numbers):
        return [], 0.0
    storage_path = Path(storage_dir).resolve()
    storage_path.mkdir(parents=True, exist_ok=True)
    try:
        model = ModelManager().get_sentence_transformer()
        if not model:
            raise HTTPException(status_code=500, detail="Failed to load embedding model")
        
        total_tokens = sum(len(c.split()) for c in chunks)
        start_time = time.time()
        embeddings = await embed_with_timeout(chunks, model)
        embed_time = time.time() - start_time
        tps = total_tokens / embed_time if embed_time > 0 else 0.0

        index_path = storage_path / index_file
        meta_path = storage_path / meta_file
        migrate_root_files(storage_dir)

        # Use cached FAISS index if available
        if FAISS_CACHE["path"] == str(index_path) and FAISS_CACHE["index"] is not None:
            index = FAISS_CACHE["index"]
            metadata_store = FAISS_CACHE["metadata"]
        else:
            index = faiss.read_index(str(index_path)) if index_path.exists() else faiss.IndexFlatIP(embeddings.shape[1])
            metadata_store = pickle.load(open(meta_path, 'rb')) if meta_path.exists() else {"ids": [], "metadata": []}
            FAISS_CACHE.update({"index": index, "metadata": metadata_store, "path": str(index_path)})

        index.add(embeddings)
        chunk_ids = [str(uuid.uuid4()) for _ in chunks]
        metadata = [
            {
                "file_name": file_name,
                "page_number": page,
                "chunk_number": i + 1,
                "content": chunk[:500],
                "created_at": datetime.utcnow().isoformat(),
                "content_hash": hash(chunk)
            }
            for i, (chunk, page) in enumerate(zip(chunks, chunk_page_numbers))
        ]
        existing_hashes = {m["content_hash"] for m in metadata_store["metadata"]}
        new_ids = []
        new_metadata = []
        for cid, meta in zip(chunk_ids, metadata):
            if meta["content_hash"] not in existing_hashes:
                new_ids.append(cid)
                new_metadata.append(meta)
        metadata_store["ids"].extend(new_ids)
        metadata_store["metadata"].extend(new_metadata)

        faiss.write_index(index, str(index_path))
        with open(meta_path, 'wb') as f:
            pickle.dump(metadata_store, f)
        return new_ids, tps
    except Exception as e:
        logger.error(f"Embedding/storage error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

def translate_text(text: str, target_lang: str, source_lang: Optional[str] = None) -> str:
    """Translate text to target_lang."""
    if not text or len(text.strip()) < 3:  # Skip short texts
        return text
    target_lang = target_lang.lower()
    if target_lang not in ["en", "ar"]:
        target_lang = "en"
    if not source_lang:
        try:
            source_lang = detect(text)
        except Exception:
            source_lang = "en"
    if source_lang == target_lang:
        return text
    try:
        intermediate_text = text
        if target_lang == "ar" and source_lang != "en":
            intermediate_text = translate_text(text, "en", source_lang)
            source_lang = "en"
        translator = ModelManager().get_translation(target_lang)
        if not translator:
            return text
        tokenizer, model = translator
        inputs = tokenizer(intermediate_text, return_tensors="pt", padding=True, truncation=True, max_length=512)
        translated = model.generate(**inputs)
        return tokenizer.batch_decode(translated, skip_special_tokens=True)[0]
    except Exception as e:
        logger.error(f"Translation error: {e}")
        return text

def summarize_text(text: str, max_length: int = 150, min_length: int = 50) -> str:
    """Summarize text using BART."""
    if not text:
        return text
    summarizer = ModelManager().get_summarizer()
    if not summarizer:
        return text
    try:
        tokenizer, model = summarizer
        inputs = tokenizer(text, return_tensors="pt", max_length=1024, truncation=True)
        summary_ids = model.generate(
            inputs["input_ids"],
            max_length=max_length,
            min_length=min_length,
            length_penalty=1.0,
            num_beams=4,
            early_stopping=True
        )
        return tokenizer.decode(summary_ids[0], skip_special_tokens=True)
    except Exception as e:
        logger.error(f"Summarization error: {e}")
        return text

def compute_rouge_scores(reference: str, hypothesis: str) -> Dict[str, float]:
    """Compute ROUGE scores."""
    try:
        scorer = rouge_scorer.RougeScorer(['rouge1', 'rouge2', 'rougeL'], use_stemmer=True)
        scores = scorer.score(reference, hypothesis)
        return {k: v.fmeasure for k, v in scores.items()}
    except Exception:
        return {"rouge1": 0.0, "rouge2": 0.0, "rougeL": 0.0}

def preprocess_query(query: str, target_lang: str = "en") -> Tuple[str, List[str]]:
    """Preprocess query with translation and NLP."""
    try:
        source_lang = detect(query)
        working_query = translate_text(query, target_lang, source_lang)
        if target_lang == "en":
            nlp = ModelManager().get_spacy()
            if not nlp:
                return working_query, []
            doc = nlp(working_query.lower())
            entities = list({token.text for token in doc if token.pos_ in ("NOUN", "PROPN") and not token.is_stop})
            normalized = " ".join(
                token.text for token in doc if token.pos_ in ("NOUN", "VERB", "ADJ", "PROPN") and not token.is_stop
            ).strip() or working_query
        else:
            entities = []
            normalized = working_query.strip()
        return normalized, entities
    except Exception as e:
        logger.error(f"Query preprocessing error: {e}")
        return query, []

async def ollama_generate_with_timeout(prompt: str, model: str = "llama3.2", timeout: float = 180.0, retries: int = 2) -> Tuple[str, float]:
    """Generate response with Ollama."""
    prompt_tokens = len(prompt.split())
    attempt = 0
    while attempt <= retries:
        try:
            available_models = []
            response = ollama.list()
            if isinstance(response, dict) and 'models' in response:
                for m in response['models']:
                    if hasattr(m, 'model') and m.model:
                        available_models.append(m.model)
                    elif isinstance(m, dict) and ('name' in m or 'model' in m):
                        available_models.append(m.get('name', m.get('model')))
            if not available_models:
                try:
                    ollama.show(model)
                    available_models.append(model)
                except Exception:
                    pass
            if not available_models:
                raise HTTPException(status_code=503, detail="No models available in Ollama")
            if model not in available_models:
                raise HTTPException(status_code=503, detail=f"Model '{model}' not found")
            
            start_time = time.time()
            response = await asyncio.wait_for(
                asyncio.get_event_loop().run_in_executor(None, lambda: ollama.generate(model=model, prompt=prompt)),
                timeout=timeout
            )
            generation_time = time.time() - start_time
            text = response.get("response", "").strip()
            tps = prompt_tokens / generation_time if generation_time > 0 else 0.0
            return text, tps
        except asyncio.TimeoutError:
            attempt += 1
            if attempt > retries:
                raise HTTPException(status_code=504, detail="Ollama timeout")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Ollama error: {str(e)}")
    return "", 0.0

@app.post("/ingest")
async def ingest_file(file: UploadFile = File(...)):
    """Ingest and process uploaded file."""
    start_time = datetime.now()
    ext = Path(file.filename).suffix.lower().lstrip(".")
    if ext not in ["pdf", "docx", "xlsx"]:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name
    
    try:
        text, page_numbers, error_report = read_file_content(tmp_path)
        if not text:
            raise HTTPException(status_code=400, detail="No content extracted")
        
        chunks, chunk_page_numbers = chunk_text((text, page_numbers))
        if not chunks:
            raise HTTPException(status_code=400, detail="No chunks generated")
        
        chunk_ids, tps = await embed_and_store_chunks(chunks, file.filename, chunk_page_numbers)
        if not chunk_ids:
            raise HTTPException(status_code=500, detail="Failed to process chunks")
        
        processing_time = (datetime.now() - start_time).total_seconds()
        return {
            "status": "success",
            "file_processed": file.filename,
            "chunks_stored": len(chunk_ids),
            "processing_time_seconds": processing_time,
            "embedding_tokens_per_second": tps,
            "metadata_storage_path": str(Path(os.getenv("STORAGE_DIR", "data")) / "faiss_metadata.pkl"),
            "errors": error_report["errors"]
        }
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

@app.post("/query")
async def query_chatbot(request: QueryRequest):
    """Process query with translation, NLP, summarization, and Ollama."""
    start_time = datetime.now()
    session_id = request.session_id or str(uuid.uuid4())
    cleanup_sessions()
    
    target_lang = request.target_language.lower()
    if target_lang not in ["en", "ar"]:
        target_lang = "en"
    
    summarize_mode = request.summarize.lower() if request.summarize else None
    if summarize_mode not in [None, "context", "answer"]:
        summarize_mode = None
    
    normalized_query, entities = preprocess_query(request.query, target_lang)
    model = ModelManager().get_sentence_transformer()
    if not model:
        raise HTTPException(status_code=500, detail="Failed to load embedding model")
    
    query_tokens = len(normalized_query.split())
    query_text = f"{normalized_query} {' '.join(entities)}" if entities else normalized_query
    embed_start = time.time()
    query_embedding = await embed_with_timeout([query_text], model, timeout=30.0)
    embed_time = time.time() - embed_start
    query_embedding_tps = query_tokens / embed_time if embed_time > 0 else 0.0
    
    storage_path = Path(os.getenv("STORAGE_DIR", "data")).resolve()
    if not storage_path.exists():
        raise HTTPException(status_code=500, detail="Storage directory missing")
    
    migrate_root_files(storage_path)
    index_path = storage_path / "faiss_index.bin"
    meta_path = storage_path / "faiss_metadata.pkl"
    
    if not index_path.exists():
        raise HTTPException(status_code=404, detail="FAISS index not found")
    
    if FAISS_CACHE["path"] == str(index_path) and FAISS_CACHE["index"] is not None:
        index = FAISS_CACHE["index"]
        metadata_store = FAISS_CACHE["metadata"]
    else:
        index = faiss.read_index(str(index_path))
        metadata_store = pickle.load(open(meta_path, 'rb')) if meta_path.exists() else {"ids": [], "metadata": []}
        FAISS_CACHE.update({"index": index, "metadata": metadata_store, "path": str(index_path)})
    
    top_k = min(request.top_k, index.ntotal)
    distances, indices = index.search(query_embedding, top_k)
    chunk_ids = metadata_store["ids"]
    metadatas = metadata_store["metadata"]
    
    retrieved_chunks = []
    sources = []
    for idx in indices[0]:
        if idx < len(chunk_ids):
            meta = metadatas[idx]
            retrieved_chunks.append(meta["content"][:1000])
            sources.append({
                "file_name": meta["file_name"],
                "page_number": meta["page_number"],
                "chunk_number": meta["chunk_number"],
                "created_at": meta["created_at"]
            })
    
    if not retrieved_chunks:
        raise HTTPException(status_code=404, detail="No relevant content found")
    
    context = "\n\n".join(retrieved_chunks)
    rouge_scores = {}
    if summarize_mode == "context":
        original_context = context
        context = summarize_text(context, max_length=200, min_length=50)
        rouge_scores = compute_rouge_scores(original_context, context)
    
    history = SESSION_HISTORY.get(session_id, [])
    history_text = "\n\nPrevious Conversation:\n" + "\n".join(
        [f"Q: {q}\nA: {a}" for q, a, _ in history[-HISTORY_LIMIT:]]
    ) if history else ""
    
    entities_text = f"Key Topics: {', '.join(entities)}" if entities else ""
    prompt = (
        "You are a helpful assistant. Use the provided context, key topics, and previous conversation to answer the query concisely and accurately. "
        f"If the context or history lacks relevant information, provide a general response.\n\n"
        f"{history_text}\n\n{entities_text}\n\nQuery: {normalized_query}\n\nContext:\n{context}"
    )
    
    answer, ollama_tps = await ollama_generate_with_timeout(prompt)
    if summarize_mode == "answer":
        original_answer = answer
        answer = summarize_text(answer, max_length=100, min_length=30)
        rouge_scores = compute_rouge_scores(original_answer, answer)
    
    source_lang = detect(request.query)
    response_lang = target_lang if target_lang != source_lang else source_lang
    translated_answer = translate_text(answer, response_lang, "en" if target_lang == "en" else "ar")
    
    SESSION_HISTORY.setdefault(session_id, []).append((request.query, translated_answer, datetime.utcnow()))
    if len(SESSION_HISTORY[session_id]) > HISTORY_LIMIT:
        SESSION_HISTORY[session_id].pop(0)
    
    processing_time = (datetime.now() - start_time).total_seconds()
    response = {
        "status": "success",
        "query": request.query,
        "answer": translated_answer,
        "session_id": session_id,
        "entities_detected": entities,
        "source_language": source_lang,
        "response_language": response_lang,
        "sources": sources,
        "processing_time_seconds": processing_time,
        "query_embedding_tps": query_embedding_tps,
        "ollama_generation_tps": ollama_tps
    }
    if summarize_mode:
        response["rouge_scores"] = rouge_scores
    return response